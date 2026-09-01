import asyncio
import tempfile
import unittest
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document
from openpyxl import load_workbook

import lqa_pipeline as pipeline


class _NeverReturningModels:
    async def generate_content(self, **kwargs):
        await asyncio.Event().wait()


class _NeverReturningClient:
    def __init__(self):
        self.aio = type("Aio", (), {"models": _NeverReturningModels()})()


class TicketPipelineTests(unittest.TestCase):
    def test_project_instructions_are_rendered_for_gr_and_pt(self):
        config = pipeline.load_config()
        prompts_dir = Path(config["paths"]["prompts"])
        project_text = (prompts_dir / config["file_names"]["project_instructions"]).read_text(
            encoding="utf-8"
        )
        templates = [
            (prompts_dir / config["file_names"][key]).read_text(encoding="utf-8")
            for key in ("system_prompt", "system_prompt_2")
        ]

        for lang_code in ("GR", "PT"):
            for template in templates:
                rendered = pipeline.render_system_prompt(
                    template,
                    lang_code,
                    f"{lang_code} locale rules",
                    project_text,
                )
                self.assertIn("Kalithèa", rendered)
                self.assertIn("Argásion", rendered)
                self.assertNotIn("$project_instructions", rendered)

    def test_language_filter_limits_discovered_files(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_dir = Path(temp_dir)
            (input_dir / "2026-2027-gr-car-hire-city-06.docx").touch()
            (input_dir / "2026-2027-pt-car-hire-city-06.docx").touch()

            files_by_lang = pipeline.discover_input_files(
                input_dir,
                "docx",
                {"GR": "el-GR", "PT": "pt-PT"},
                "GR",
            )

        self.assertEqual(list(files_by_lang), ["GR"])
        self.assertEqual([p.name for p in files_by_lang["GR"]], ["2026-2027-gr-car-hire-city-06.docx"])

    def test_gr_project_filter_removes_transliteration_but_keeps_greek_typo(self):
        original = "Στην Kalithèa. Η τιμή στην Πετρα είναι 10 €."
        findings = [
            {
                "text": "Στην Kalithèa.",
                "issue": "The place name Kalithèa is an incorrect transliteration.",
                "updated_sentence": "Στην Καλλιθέα.",
            },
            {
                "text": "Η τιμή στην Πετρα είναι 10 €.",
                "issue": "The Greek place name Πετρα is misspelled.",
                "updated_sentence": "Η τιμή στην Πέτρα είναι 10 €.",
            },
        ]

        kept, corrected, removed = pipeline.apply_project_finding_rules(
            "GR",
            original,
            findings,
            "Στην Καλλιθέα. Η τιμή στην Πέτρα είναι 10 €.",
        )

        self.assertEqual([f["issue"] for f in kept], [findings[1]["issue"]])
        self.assertEqual(corrected, "Στην Kalithèa. Η τιμή στην Πέτρα είναι 10 €.")
        self.assertEqual([f["issue"] for f in removed], [findings[0]["issue"]])

    def test_agent_audit_persists_agent_counts(self):
        rows = [{"section_id": "gr-1", "source_file": "/input/gr.docx"}]
        agent1 = [{"review_findings": [{"issue": "draft"}, {"issue": "draft 2"}]}]
        agent2 = [{"review_findings": [{"issue": "kept"}]}]

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "agent-audit.json"
            pipeline.write_agent_audit(output_path, "GR", rows, agent1, agent2)
            payload = yaml.safe_load(output_path.read_text(encoding="utf-8"))

        self.assertEqual(payload["agent1_finding_count"], 2)
        self.assertEqual(payload["agent2_finding_count"], 1)
        self.assertEqual(payload["sections"][0]["section_id"], "gr-1")

    def test_agent2_operational_error_does_not_reduce_mqm(self):
        issue = {
            "issue": "Agent 2 error: request timed out",
            "type_of_issue": "Accuracy",
            "severity": "Major",
            "score_deduction": 3,
        }

        metrics = pipeline.aggregate_metrics([issue], ewt=100)

        self.assertEqual(metrics["apt"], 0)
        self.assertEqual(metrics["mqm"], 100.0)
        self.assertEqual(metrics["issues"][0]["score_deduction"], 0)

    def test_thai_edited_docx_has_explicit_font_and_language_metadata(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            source = temp / "source.docx"
            output = temp / "edited.docx"
            doc = Document()
            doc.add_heading("https://example.test", level=1)
            doc.add_heading("หัวข้อทดสอบ", level=2)
            doc.add_paragraph("ข้อความภาษาไทย")
            doc.save(source)
            rows = [{
                "section_id": "th-1",
                "source_file": str(source),
                "url_path": "https://example.test",
                "heading": "หัวข้อทดสอบ",
                "content": "ข้อความภาษาไทย",
            }]

            pipeline.write_edited_docx(source, rows, {}, output, lang_code="TH")

            with ZipFile(output) as docx_zip:
                xml = docx_zip.read("word/document.xml").decode("utf-8")

        self.assertIn('w:eastAsia="Tahoma"', xml)
        self.assertIn('w:cs="Tahoma"', xml)
        self.assertIn('w:eastAsia="th-TH"', xml)
        self.assertIn('w:bidi="th-TH"', xml)
        self.assertIn("ข้อความภาษาไทย", xml)

    def test_language_filter_is_case_insensitive(self):
        self.assertEqual(
            pipeline._normalize_requested_language("gr", {"CZ": "cs-CZ", "GR": "el-GR"}),
            "GR",
        )

    def test_language_filter_rejects_unknown_code(self):
        with self.assertRaisesRegex(ValueError, "Unknown language code"):
            pipeline._normalize_requested_language("XX", {"CZ": "cs-CZ", "GR": "el-GR"})

    def test_three_parallel_runs_are_capped_at_thirty_combined_calls(self):
        config = pipeline.load_config()
        self.assertLessEqual(config["batching"]["max_batches_per_window"] * 3, 30)

    def test_all_ticket_docx_files_map_to_their_runtime_locale(self):
        config = pipeline.load_config()
        archive = Path("03_Car Hire in City-20260831T145126Z-1-001.zip")

        with ZipFile(archive) as zip_file:
            docx_paths = [Path(name) for name in zip_file.namelist() if name.endswith(".docx")]

        counts = Counter(
            pipeline._infer_lang_code_for_input_file(path, config["lang_map"])
            for path in docx_paths
        )

        self.assertEqual(
            counts,
            {"CZ": 5, "GR": 5, "HU": 5, "PL": 5, "PT": 5, "SA": 5, "TH": 5, "TR": 5},
        )

    def test_detailed_findings_report_shows_source_filename(self):
        issue = {
            "source_file": "/ticket/input/example.docx",
            "candidate": "https://example.test/page",
            "text": "Original text",
            "issue": "Objective error",
            "type_of_issue": "Grammar",
            "severity": "Minor",
            "score_deduction": 1,
            "solution": "Apply the correction.",
            "updated_sentence": "Corrected text",
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir)
            pipeline.write_language_report(
                lang_code="PT",
                evaluated_scope_label="test",
                evaluated_text="Original text",
                issues_with_url=[issue],
                output_dir=output_dir,
                save_raw_json=False,
            )
            report_path = next(output_dir.glob("*.xlsx"))
            workbook = load_workbook(report_path, read_only=True, data_only=True)
            rows = list(workbook["Detailed Findings"].iter_rows(values_only=True))
            workbook.close()

        self.assertEqual(rows[0][0], "Source File")
        self.assertEqual(rows[1][0], "example.docx")


class GeminiTimeoutTests(unittest.IsolatedAsyncioTestCase):
    async def test_hung_gemini_request_is_bounded_by_pipeline_timeout(self):
        started = asyncio.get_running_loop().time()
        original_timeout = getattr(pipeline, "GEMINI_REQUEST_TIMEOUT_SECONDS", None)
        pipeline.GEMINI_REQUEST_TIMEOUT_SECONDS = 0.01
        try:
            with self.assertRaises(asyncio.TimeoutError):
                await asyncio.wait_for(
                    pipeline._call_gemini_once.__wrapped__(
                        _NeverReturningClient(),
                        "gemini-2.5-pro",
                        1,
                        {"system": "test", "prompt": "test"},
                    ),
                    timeout=0.2,
                )
        finally:
            if original_timeout is None:
                del pipeline.GEMINI_REQUEST_TIMEOUT_SECONDS
            else:
                pipeline.GEMINI_REQUEST_TIMEOUT_SECONDS = original_timeout

        self.assertLess(asyncio.get_running_loop().time() - started, 0.1)


if __name__ == "__main__":
    unittest.main()
