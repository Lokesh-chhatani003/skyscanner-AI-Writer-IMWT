# lqa_pipeline.py
# ------------------------------------------------------------------------
# Skyscanner AI-LQA pipeline (DOCX-first, windowed per-language, three-agent)
#   - Agent 1: Draft per-section findings
#   - Agent 2: Validate/finalize per-section findings (+ corrected_content)
#   - Agent 3: Consolidate language-level Summary & Action Plan
# Deliverables: scorecard XLSX/JSON + edited DOCX per input file
# ------------------------------------------------------------------------

from __future__ import annotations

import argparse
import asyncio
import csv
import json
import logging
import os
import re
from collections import defaultdict, Counter
from datetime import datetime
from html import escape as html_escape
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai
from google.genai import types
from tenacity import retry, stop_after_attempt, wait_random_exponential
from zoneinfo import ZoneInfo

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo

# ────────────────────────── LOGGING SETUP ──────────────────────────────────
# Console logging immediately; file handler is added inside main() once we know output dir
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)
GEMINI_REQUEST_TIMEOUT_SECONDS = 600


# ─────────────────────── GEMINI HELPER FUNCTIONS ──────────────────────────
def extract_json(raw_text: str) -> Dict[str, Any]:
    """Safely extract a JSON object from a string that might contain extra text."""
    try:
        json_start = raw_text.find("{")
        if json_start == -1:
            raise ValueError("No JSON object found in the response.")
        json_end = raw_text.rfind("}")
        if json_end == -1:
            raise ValueError("JSON object is incomplete.")
        json_str = raw_text[json_start : json_end + 1]
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        logging.error(f"Failed to decode JSON: {e}")
        logging.error(f"Raw text was: {raw_text}")
        raise
    except ValueError as e:
        logging.error(f"Error extracting JSON: {e}")
        logging.error(f"Raw text was: {raw_text}")
        raise


@retry(stop=stop_after_attempt(3), wait=wait_random_exponential(multiplier=1, max=60))
async def _call_gemini_once(
    client: genai.Client,
    llm_model: str,
    temp: float,
    parts: Dict[str, str],
) -> Dict[str, Any]:
    """
    One Gemini API call. Tenacity handles retries externally.
    """
    config_params: Dict[str, Any] = {
        "system_instruction": parts["system"],
        "temperature": temp,
        "max_output_tokens": 32_000,
        "response_mime_type": "text/plain",
    }
    # thinking_level is Gemini 3 only; 2.x models reject it.
    if "gemini-3" in llm_model:
        config_params["thinking_config"] = types.ThinkingConfig(thinking_level="high")

    response = await asyncio.wait_for(
        client.aio.models.generate_content(
            model=llm_model,
            contents=parts["prompt"],
            config=types.GenerateContentConfig(**config_params),
        ),
        timeout=GEMINI_REQUEST_TIMEOUT_SECONDS,
    )
    if not response.text:
        raise ValueError("Empty response from Gemini")

    raw = response.text
    return extract_json(raw)


async def get_gemini_lqa_batch(
    api_key: str,
    llm_model: str,
    temp: float,
    prompt_parts_list: List[Dict[str, str]],
) -> List[Dict[str, Any]]:
    """Run a list of Gemini requests concurrently and return results aligned to inputs."""
    client = genai.Client(api_key=api_key)

    async def _safe_call(parts):
        try:
            return await _call_gemini_once(client, llm_model, temp, parts)
        except Exception as exc:  # noqa: BLE001
            # Tenacity wraps the real failure; unwrap so the log shows the API message.
            cause = getattr(exc, "last_attempt", None)
            if cause is not None:
                try:
                    exc = cause.exception() or exc
                except Exception:  # noqa: BLE001
                    pass
            logger.error("Gemini task failed (model=%s): %s", llm_model, exc)
            return {"error": f"{type(exc).__name__}: {exc}"}

    return await asyncio.gather(*(_safe_call(p) for p in prompt_parts_list))


# ──────────────────────── CONFIG / DATA LOADING ────────────────────────────
def load_config(config_path: str = "config.yaml") -> Dict[str, Any]:
    logging.info(f"Loading configuration from {config_path}")
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def resolve_api_key(config: Dict[str, Any]) -> str:
    """
    Resolve Gemini API key with priority:
    1) GEMINI_API_KEY environment variable
    2) config.yaml -> api_key
    """
    env_key = os.getenv("GEMINI_API_KEY", "").strip()
    if env_key:
        logging.info("Using Gemini API key from environment variable GEMINI_API_KEY.")
        return env_key

    config_key_raw = config.get("api_key", "") if isinstance(config, dict) else ""
    config_key = str(config_key_raw).strip() if config_key_raw is not None else ""
    placeholder_values = {
        "",
        "<YOUR_GEMINI_API_KEY>",
        "YOUR_GEMINI_API_KEY",
        "__SET_ME__",
        "REPLACE_ME",
    }
    if config_key and config_key not in placeholder_values:
        logging.info("Using Gemini API key from config.yaml (api_key).")
        return config_key

    raise ValueError(
        "Gemini API key is missing. Set GEMINI_API_KEY (recommended), or set a non-placeholder "
        "value in config.yaml under 'api_key'."
    )


def _cell_text(value: Any) -> str:
    """Normalize a worksheet cell to trimmed text ('' for blanks)."""
    if value is None:
        return ""
    return str(value).strip()


def load_glossary(
    ref_path: Path, filename: str, lang_map: Dict[str, str]
) -> Dict[str, str]:
    """
    Load Skyscanner_Glossary_v2-style Multilingual Glossary and format
    compact en-GB ↔ locale term lines for each short language code.
    """
    glossary_path = ref_path / filename
    logging.info(f"Loading and formatting glossary from {glossary_path}")
    glossary_cache: Dict[str, str] = {code: "N/A" for code in (lang_map or {})}

    try:
        wb = load_workbook(glossary_path, read_only=True, data_only=True)
    except FileNotFoundError:
        logging.error(f"Glossary file not found at {glossary_path}.")
        return glossary_cache
    except Exception as e:
        logging.error(f"Glossary load error: {e}", exc_info=True)
        return glossary_cache

    try:
        sheet_name = "Multilingual Glossary"
        if sheet_name not in wb.sheetnames:
            sheet_name = wb.sheetnames[0]
            logging.warning("Sheet 'Multilingual Glossary' not found; using '%s'.", sheet_name)
        ws = wb[sheet_name]
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
    except Exception as e:
        logging.error(f"Glossary load error: {e}", exc_info=True)
        return glossary_cache
    finally:
        try:
            wb.close()
        except Exception:  # noqa: BLE001
            pass

    if len(rows) < 2:
        logging.error("Glossary workbook is empty.")
        return glossary_cache

    headers = [_cell_text(h) for h in rows[0]]
    data = rows[1:]

    def _col_index(name: str) -> Optional[int]:
        target = name.strip().lower()
        for i, header in enumerate(headers):
            if header.lower() == target:
                return i
        return None

    def_idx = _col_index("Definition")
    en_idx = _col_index("en-GB")

    for short_code, long_code in (lang_map or {}).items():
        loc_idx = _col_index(str(long_code))
        if loc_idx is None:
            logging.warning("Glossary: locale column %s not found for %s", long_code, short_code)
            continue

        notes_idx = (
            loc_idx + 1
            if loc_idx + 1 < len(headers) and headers[loc_idx + 1].lower() == "notes"
            else None
        )
        example_idx = (
            loc_idx + 2
            if loc_idx + 2 < len(headers) and "example" in headers[loc_idx + 2].lower()
            else None
        )

        def _value(row: List[Any], idx: Optional[int]) -> str:
            if idx is None or idx >= len(row):
                return ""
            return _cell_text(row[idx])

        lines: List[str] = ["### Skyscanner Multilingual Glossary ###"]
        for row in data:
            loc_s = _value(row, loc_idx)
            if not loc_s:
                continue
            en_s = _value(row, en_idx)
            def_s = _value(row, def_idx)
            notes_s = _value(row, notes_idx)
            example_s = _value(row, example_idx)

            piece = f"- EN: {en_s} → {long_code}: {loc_s}" if en_s else f"- {long_code}: {loc_s}"
            if def_s:
                piece += f" | Definition: {def_s}"
            if notes_s:
                piece += f" | Notes: {notes_s}"
            if example_s:
                piece += f" | Example: {example_s}"
            lines.append(piece)

        glossary_cache[short_code] = "\n".join(lines) if len(lines) > 1 else "N/A"
        logging.info(
            "Glossary loaded for %s (%s): %d terms",
            short_code,
            long_code,
            max(0, len(lines) - 1),
        )

    return glossary_cache


# ───────────────────────── HTML → sections (markdown-like) ────────────────
try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError as e:
    raise ImportError("Please install: pip install beautifulsoup4 lxml html5lib") from e

def _pick_parser() -> str:
    for p in ("html5lib", "lxml", "html.parser"):
        try:
            BeautifulSoup("<div></div>", p)
            return p
        except Exception:
            continue
    return "html.parser"

_PARSER = _pick_parser()

# NBSP-aware trim: do NOT strip NBSP (U+00A0) or NNBSP (U+202F)
_WS_EDGE_RE = re.compile(r"^[ \t\r\n]+|[ \t\r\n]+$")
def _strip_layout(s: str) -> str:
    return _WS_EDGE_RE.sub("", s)

def _normalize_ws(s: str) -> str:
    # collapse only normal spaces/tabs; keep NBSP (\u00A0) / NNBSP (\u202F) intact
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return _strip_layout(s)

def _heading_prefix(tag_name: str) -> str:
    level = int(tag_name[1]) if tag_name and tag_name.startswith("h") and tag_name[1].isdigit() else 3
    return "#" * max(1, min(level, 6))

# Exact inline text (no invented joiner spaces). Preserves author-provided spacing incl. NBSP.
def _text_preserve_inline(node: Tag) -> str:
    parts: List[str] = []
    for child in node.descendants:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif isinstance(child, Tag) and child.name.lower() == "br":
            parts.append("\n")
    return "".join(parts)

def _get_immediate_text(node: Tag) -> str:
    """
    Extract immediate inline text from `node` without inventing spaces at tag boundaries.
    Preserves author-provided spacing exactly (incl. NBSP/NNBSP).
    """
    pieces: List[str] = []
    for child in node.children:
        if isinstance(child, NavigableString):
            pieces.append(str(child))
        elif isinstance(child, Tag):
            if child.name in ("ul", "ol"):
                continue  # let _render_list handle nested lists
            pieces.append(_text_preserve_inline(child))
    return _strip_layout("".join(pieces))

def _render_list(node: Tag, depth: int = 0) -> str:
    if not isinstance(node, Tag) or node.name not in ("ul", "ol"):
        return ""
    lines: List[str] = []
    last_li: Optional[Tag] = None

    for child in node.children:
        if not isinstance(child, Tag):
            continue
        name = child.name.lower()

        if name == "li":
            last_li = child
            main = _get_immediate_text(child)
            if main:
                lines.append(("  " * depth) + "- " + main)

            # Properly nested lists inside this <li>
            for sub in child.find_all(["ul", "ol"], recursive=False):
                nested_txt = _render_list(sub, depth + 1)
                if nested_txt:
                    lines.append(nested_txt)

        elif name in ("ul", "ol"):
            # Invalid-but-common markup: a sublist placed as a sibling after the <li>.
            if last_li is not None:
                nested_txt = _render_list(child, depth + 1)  # attach to prior <li>
            else:
                nested_txt = _render_list(child, depth)      # fallback
            if nested_txt:
                lines.append(nested_txt)

    return "\n".join(lines)


def _is_leaf_text_div(node: Tag) -> bool:
    """
    Some component libraries render paragraph-like content in <div> nodes (e.g., accordions).
    Treat a <div> as a paragraph only when it's a true text leaf to avoid duplicating
    content from layout wrappers.
    """
    if not isinstance(node, Tag) or node.name != "div":
        return False

    # If this div contains other block-ish elements (incl. nested divs), treat it as layout.
    if node.find(
        [
            "div",
            "p",
            "ul",
            "ol",
            "li",
            "table",
            "section",
            "article",
            "header",
            "footer",
            "nav",
            "aside",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
        ]
    ):
        return False

    txt = _strip_layout(_text_preserve_inline(node))
    if not txt:
        return False

    # Avoid ultra-short UI fragments (e.g., separators like "/" or "Caption").
    words = [w for w in re.split(r"\s+", txt) if w]
    return len(words) >= 3

def _render_block(node: Tag) -> str:
    out: List[str] = []
    for el in node.descendants:
        if not isinstance(el, Tag):
            continue
        name = el.name.lower()
        if name in ("script", "style", "noscript"):
            continue
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            txt = _strip_layout(_text_preserve_inline(el))
            if txt:
                out.append(f"{_heading_prefix(name)} {txt}")
        elif name == "p":
            txt = _strip_layout(_text_preserve_inline(el))
            if txt:
                out.append(txt)
        elif name in ("ul", "ol"):
            # Skip nested lists - only render top-level lists
            # Check if this list is nested inside another list (parent ul/ol)
            if el.find_parent(["ul", "ol"]) is not None:
                continue
            rendered = _render_list(el, depth=0)
            if rendered:
                out.append(rendered)
        elif name == "div" and _is_leaf_text_div(el):
            txt = _strip_layout(_text_preserve_inline(el))
            if txt:
                out.append(txt)
        if name == "br":
            out.append("")
    return _normalize_ws("\n\n".join(out))

def _find_two_col_cells(section_soup: BeautifulSoup) -> Optional[Tuple[Tag, Tag]]:
    for tbl in section_soup.find_all("table"):
        rows = tbl.find_all("tr", recursive=True)
        if not rows:
            continue
        row = next((r for r in rows if r.find_all(["td","th"], recursive=False)), None)
        if not row:
            continue
        tds = row.find_all("td", recursive=False)
        if len(tds) >= 2:
            return (tds[0], tds[1])
    return None

def _choose_content_and_meta(section_soup: BeautifulSoup) -> Tuple[Tag, Optional[Tag]]:
    pair = _find_two_col_cells(section_soup)
    if pair:
        left, right = pair
        left_text = left.get_text(" ", strip=True).lower()
        if "candidate" in left_text or "urlpath" in left_text:
            return right, left
        def score(td: Tag) -> int:
            return len(td.find_all(["h1","h2","h3","p","ul","ol","li"]))
        return (right, left) if score(right) >= score(left) else (left, right)

    # Newer "Destination Hub" exports (and some other templates) do not use the legacy
    # two-column metadata table layout, but do include a single content root.
    deliverable = section_soup.select_one(".tm-deliverable")
    if deliverable:
        return (deliverable, None)

    for tag_name in ("main", "article"):
        root = section_soup.find(tag_name)
        if root:
            return (root, None)

    return (section_soup.body or section_soup, None)

def _first_heading_text(container: Tag) -> Optional[str]:
    for lvl in ("h1","h2","h3","h4","h5","h6"):
        tag = container.find(lvl)
        if tag:
            t = tag.get_text(" ", strip=True)
            if t:
                return t
    return None

_URLPATH_JSON_RE = re.compile(r"(?i)\"urlPath\"\s*:\s*\"([^\"]+)\"")

def _extract_url_path_from_text(text: str) -> Optional[str]:
    if not text:
        return None
    m = _URLPATH_JSON_RE.search(text)
    if m:
        return m.group(1).strip()
    m = re.search(r"url\s*path\s*[:=]\s*(/[^\s\"'>]+)", text, flags=re.I)
    if m:
        return m.group(1).strip()
    return None

def _extract_url_path(meta_td: Optional[Tag], html_fallback: Optional[str] = None) -> Optional[str]:
    if meta_td:
        # Prefer explicit hrefs
        for a in meta_td.find_all("a", href=True):
            href = a["href"].strip()
            if re.search(r"/[^\s\"'>]+\.html?$", href, flags=re.I):
                return href
        # Try JSON or labeled urlPath in meta text
        txt = meta_td.get_text(" ", strip=True)
        found = _extract_url_path_from_text(txt)
        if found:
            return found
        # Last resort: any /...html in meta text
        m2 = re.search(r"(/[^\s\"'>]+\.html?)", txt, flags=re.I)
        if m2:
            return m2.group(1)

    if html_fallback:
        return _extract_url_path_from_text(html_fallback)
    return None


_SAVED_FROM_URL_RE = re.compile(r"(?i)saved from url=\(\d+\)(https?://[^\s]+)")

def _fallback_candidate_path_from_html(html: str, html_path: str | Path) -> str:
    """
    Destination-Hub exports often contain a leading comment like:
      <!-- saved from url=(0107)https://example.com/path/file.html -->
    Use it as a stable candidate identifier when the legacy metadata table is absent.
    """
    m = _SAVED_FROM_URL_RE.search(html or "")
    if m:
        raw_url = m.group(1)
        try:
            parsed = urlparse(raw_url)
            if parsed.scheme and parsed.netloc and parsed.path:
                return parsed.path
        except Exception:
            pass
        return raw_url

    p = Path(html_path)
    return p.name

def _fallback_topic_name_from_candidate(candidate_path: str) -> Optional[str]:
    """
    Best-effort topic extraction from a candidate path/filename, e.g.:
      /.../shinjuku_fr_041225_01.html -> "Shinjuku"
    """
    if not candidate_path:
        return None
    name = str(candidate_path).rsplit("/", 1)[-1]
    name = re.sub(r"\.html?$", "", name, flags=re.I)
    slug = name.split("_", 1)[0].strip()
    if not slug:
        return None
    slug = slug.replace("-", " ").strip()
    return slug.title() if slug else None

def _extract_topic_name(meta_td: Optional[Tag]) -> Optional[str]:
    """
    Extracts a human-readable topic name from the metadata column.
    Prefers explicit labels like "CANDIDATE NAME" or "TOPIC NAME".
    Falls back to None if not found.
    """
    if not meta_td:
        return None

    # Build a mapping of label -> value using <strong> markers as keys.
    topic_map: Dict[str, str] = {}
    for strong in meta_td.find_all("strong"):
        key = strong.get_text(" ", strip=True).strip().lower()
        val_parts: List[str] = []
        for sib in strong.next_siblings:
            if isinstance(sib, Tag) and sib.name.lower() == "br":
                break
            val_parts.append(str(sib))
        val = _strip_layout("".join(val_parts))
        if val:
            topic_map[key] = val

    for k in ("topic name", "candidate name"):
        if k in topic_map:
            return topic_map[k]

    return None

def _extract_topic_name_from_html(html: str) -> Optional[str]:
    if not html:
        return None
    try:
        soup = BeautifulSoup(html, _PARSER)
    except Exception:
        return None
    return _extract_topic_name(soup)

def _candidate_name_from_url_path(url_path: Optional[str]) -> Optional[str]:
    if not url_path:
        return None
    path = str(url_path).split("?", 1)[0].split("#", 1)[0]
    slug = path.rsplit("/", 1)[-1]
    slug = re.sub(r"\.html?$", "", slug, flags=re.I)
    slug = re.sub(r"(?i)^(cheap-)?flights?-to-", "", slug)
    slug = re.sub(r"(?i)^(car-rentals?-in|car-hire-from|car-hire-in|hotels?-in)-", "", slug)
    slug = slug.replace("-", " ").strip()
    return slug.title() if slug else None

def _clean_candidate_name(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    cleaned = str(value).strip()
    cleaned = re.sub(r"^[\s:\-–—]+", "", cleaned)
    return cleaned if cleaned else None

def _count_hr_sections(html: str) -> int:
    parts = re.split(r"(?i)<hr\b[^>]*>", html or "")
    return sum(1 for seg in parts if seg.strip())


def parse_html_sections_to_df(
    html_path: str | Path,
    drop_empty: bool = True,
    max_hr_sections: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """
    Split HTML by <hr>. For each section, pick the content column, render markdown-like text,
    and extract urlPath from metadata column. Returns: section_id, url_path, content.
    """
    html = Path(html_path).read_text(encoding="utf-8", errors="ignore")
    fallback_candidate = _fallback_candidate_path_from_html(html, html_path)
    fallback_topic_name = _fallback_topic_name_from_candidate(fallback_candidate)
    parts = re.split(r"(?i)<hr\b[^>]*>", html)
    rows: List[dict] = []

    seen = 0
    for idx, seg in enumerate(parts, start=1):
        seg_clean = seg.strip()
        if not seg_clean:
            continue
        seen += 1
        if isinstance(max_hr_sections, int) and seen > max_hr_sections:
            break
        soup = BeautifulSoup(seg_clean, _PARSER)
        content_td, meta_td = _choose_content_and_meta(soup)
        content_md = _render_block(content_td)
        if drop_empty and not content_md.strip():
            continue  # skip pure-metadata sections

        heading = _first_heading_text(content_td)
        base_id = heading if heading else f"section-{idx:03d}"
        safe_id = re.sub(r"\s+", " ", base_id).strip()
        safe_id = re.sub(r"[/\\|:*?\"<>]+", "-", safe_id)
        if len(safe_id) > 120:
            safe_id = safe_id[:120].rstrip()

        url_path = _extract_url_path(meta_td, seg_clean) or fallback_candidate
        raw_name = _extract_topic_name(meta_td) or _extract_topic_name_from_html(seg_clean)
        candidate_name = _clean_candidate_name(raw_name) or _candidate_name_from_url_path(url_path)
        topic_name = candidate_name or fallback_topic_name
        rows.append(
            {
                "section_id": safe_id,
                "url_path": url_path,
                "content": content_md,
                "topic_name": topic_name,
                "candidate_name": candidate_name,
            }
        )

    return rows


# ───────────── NEW: heading-based sub-batching inside each <hr> section ─────────────
def _sanitize_id(text: str, fallback: str) -> str:
    if not text:
        return fallback
    safe = re.sub(r"\s+", " ", text).strip()
    safe = re.sub(r"[/\\|:*?\"<>]+", "-", safe)
    return safe[:120].rstrip() if len(safe) > 120 else safe

def _extract_parent_h1_text(container: Tag) -> Optional[str]:
    h1 = container.find("h1")
    if h1:
        t = h1.get_text(" ", strip=True)
        return t if t else None
    return None

def _split_content_by_headings(
    content_td: Tag,
    prefer_level: int = 2,
    include_parent_h1: str = "prefix_all",   # "prefix_all" | "first_only" | "none"
    create_intro_batch_if_long: bool = True,
    min_intro_chars: int = 200,
    allow_fallback_to_h1: bool = False,
) -> Optional[List[Tuple[str, str]]]:
    """
    Return list of (heading_text, rendered_markdown) sub-batches for a section content cell.
    - Splits only at the configured level (prefer_level). Child headings remain in the same batch.
    - If there is a parent H1 at the top, we can prefix it to batches for context (configurable).
    - Intro text before first heading is merged into first batch unless long.
    """
    # Honor prefer_level; optionally fall back to H1 if allowed
    level = None
    if content_td.find(f"h{prefer_level}"):
        level = prefer_level
    elif allow_fallback_to_h1 and len(content_td.find_all("h1")) >= 2:
        level = 1
    else:
        return None  # no splitting for this section

    parent_h1 = _extract_parent_h1_text(content_td)

    # Work on the raw HTML string for robust slicing between same-level headings
    html = str(content_td)
    # capture <hN ...> ... </hN> including attrs, inner HTML; group(0) includes heading tag
    pat = re.compile(rf"(?is)(<h{level}\b[^>]*>)(.*?)</h{level}>")
    matches = list(pat.finditer(html))
    if not matches:
        return None

    # Intro (before the first heading at chosen level)
    intro_html = html[: matches[0].start()]
    intro_md = ""
    if intro_html and intro_html.strip():
        intro_md = _render_block(BeautifulSoup(intro_html, _PARSER))

    # Build chunks: each heading token + content until next same-level heading
    chunks: List[Tuple[str, str]] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(html)
        chunk_html = html[start:end]
        heading_text = BeautifulSoup(m.group(0), _PARSER).get_text(" ", strip=True)
        chunk_md = _render_block(BeautifulSoup(chunk_html, _PARSER))

        # Prefix with parent H1 for context if requested
        if parent_h1 and include_parent_h1 in ("prefix_all", "first_only"):
            if include_parent_h1 == "prefix_all" or (include_parent_h1 == "first_only" and i == 0):
                chunk_md = f"Parent: {parent_h1}\n\n{chunk_md}"

        chunks.append((heading_text, chunk_md))

    # Handle intro text
    if intro_md:
        if create_intro_batch_if_long and len(intro_md) >= min_intro_chars:
            intro_heading = "Intro"
            if parent_h1 and include_parent_h1 in ("prefix_all", "first_only"):
                intro_md = f"Parent: {parent_h1}\n\n{intro_md}"
            chunks.insert(0, (intro_heading, intro_md))
        else:
            htxt, md = chunks[0]
            chunks[0] = (htxt, f"{intro_md}\n\n{md}")

    return chunks

def parse_html_sections_to_df_heading(
    html_path: str | Path,
    drop_empty: bool = True,
    prefer_level: int = 2,
    include_parent_h1: str = "prefix_all",
    create_intro_batch_if_long: bool = True,
    min_intro_chars: int = 200,
    allow_fallback_to_h1: bool = False,
    max_hr_sections: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """
    Split by <hr>, then within each section's content cell, sub-split by headings:
    - Prefer the configured level (default <h2>). If not found and allow_fallback_to_h1=True and multiple <h1> exist, split by <h1>.
    - Child headings (e.g., <h3>) stay within the parent batch.
    Returns rows with: section_id, url_path, content (markdown-like).
    """
    html = Path(html_path).read_text(encoding="utf-8", errors="ignore")
    fallback_candidate = _fallback_candidate_path_from_html(html, html_path)
    fallback_topic_name = _fallback_topic_name_from_candidate(fallback_candidate)
    parts = re.split(r"(?i)<hr\b[^>]*>", html)
    rows: List[dict] = []

    seen = 0
    for sidx, seg in enumerate(parts, start=1):
        seg_clean = seg.strip()
        if not seg_clean:
            continue
        seen += 1
        if isinstance(max_hr_sections, int) and seen > max_hr_sections:
            break
        soup = BeautifulSoup(seg_clean, _PARSER)
        content_td, meta_td = _choose_content_and_meta(soup)

        # Try heading-based split for this section
        sub = _split_content_by_headings(
            content_td=content_td,
            prefer_level=prefer_level,
            include_parent_h1=include_parent_h1,
            create_intro_batch_if_long=create_intro_batch_if_long,
            min_intro_chars=min_intro_chars,
            allow_fallback_to_h1=allow_fallback_to_h1,
        )

        url_path = _extract_url_path(meta_td, seg_clean) or fallback_candidate
        raw_name = _extract_topic_name(meta_td) or _extract_topic_name_from_html(seg_clean)
        candidate_name = _clean_candidate_name(raw_name) or _candidate_name_from_url_path(url_path)
        topic_name = candidate_name or fallback_topic_name

        if not sub:
            # Fallback: single batch for this section (same as HR behavior)
            content_md = _render_block(content_td)
            if drop_empty and not content_md.strip():
                continue
            heading = _first_heading_text(content_td)
            base_id = heading if heading else f"section-{sidx:03d}"
            rows.append(
                {
                    "section_id": _sanitize_id(base_id, f"section-{sidx:03d}"),
                    "url_path": url_path,
                    "content": content_md,
                    "topic_name": topic_name,
                    "candidate_name": candidate_name,
                }
            )
            continue

        # Build sub-batches with deterministic IDs
        for hidx, (heading_text, chunk_md) in enumerate(sub, start=1):
            if drop_empty and not str(chunk_md).strip():
                continue
            label = f"s{sidx:03d}-h{hidx:02d} ▸ {heading_text}"
            rows.append(
                {
                    "section_id": _sanitize_id(label, f"s{sidx:03d}-h{hidx:02d}"),
                    "url_path": url_path,
                    "content": chunk_md,
                    "topic_name": topic_name,
                    "candidate_name": candidate_name,
                }
            )

    return rows
# ─────────── END NEW ───────────


# ──────────────────────── METRICS / AGGREGATION ────────────────────────────
WORD_RE = re.compile(
    r"[\w'-]+|[\u4e00-\u9fff]+|[\u3040-\u30ff]+|[\uac00-\ud7af]+|[\u1100-\u11ff]+|[\u3130-\u318f]+",
    re.UNICODE,
)

ALLOWED_CATEGORIES = [
    "Grammar",
    "Compliance",
    "Punctuation/Spelling/Typos",
    "Accuracy",
    "Style",
    "AI-issues",
]

SEVERITY_TO_POINTS = {"Minor": 1, "Major": 3, "Critical": 25}


def compute_ewt(text: str) -> int:
    cleaned = re.sub(r"[<>]|`{1,3}.*?`{1,3}", " ", text, flags=re.S)
    return len(WORD_RE.findall(cleaned))


def normalize_issue(issue: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(issue)
    # map severity → points
    sev = str(out.get("severity", "")).strip().title()
    pts = SEVERITY_TO_POINTS.get(sev, 0)
    out["severity"] = sev if sev in SEVERITY_TO_POINTS else "Minor" if pts == 1 else sev
    out["score_deduction"] = pts
    # normalize categories
    cat = str(out.get("type_of_issue", "")).strip()
    canon = {
        "ai-issues": "AI-issues",
        "ai issues": "AI-issues",
        "punctuation/spelling/typo": "Punctuation/Spelling/Typos",
        "punctuation/spelling/typos": "Punctuation/Spelling/Typos",
        "consistency": "Compliance",
    }
    out["type_of_issue"] = canon.get(cat.lower(), cat)
    # remove unwanted fields if any
    out.pop("question", None)
    return out


def _is_agent_operational_error(issue: Dict[str, Any]) -> bool:
    message = str(issue.get("issue", ""))
    return message.startswith(("Agent 2 error:", "Agent 2 returned an unexpected payload"))


def _issues_for_reporting_and_scoring(
    issues: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    reported = [normalize_issue(issue) for issue in issues]
    for issue in reported:
        if _is_agent_operational_error(issue):
            issue["score_deduction"] = 0
    return reported, [issue for issue in reported if not _is_agent_operational_error(issue)]


def aggregate_metrics(
    issues: List[Dict[str, Any]],
    ewt: int,
    pass_threshold: float = 90.0,
) -> Dict[str, Any]:
    reported, scorable = _issues_for_reporting_and_scoring(issues)
    apt = sum(i.get("score_deduction", 0) for i in scorable)
    pwpt = (apt / ewt) if ewt > 0 else 0.0
    mqm = max(0.0, 100.0 - (pwpt * 1000.0))
    crit_cnt = sum(1 for i in scorable if i.get("severity") == "Critical")
    status = "Pass" if (mqm >= pass_threshold and crit_cnt == 0) else "Fail"

    by_cat_points = defaultdict(int)
    for i in scorable:
        cat = i.get("type_of_issue", "")
        if cat in ALLOWED_CATEGORIES:
            by_cat_points[cat] += i.get("score_deduction", 0)

    categories = []
    for cat in ALLOWED_CATEGORIES:
        cat_apt = by_cat_points.get(cat, 0)
        cat_pwpt = (cat_apt / ewt) if ewt > 0 else 0.0
        cat_mqm = max(0.0, 100.0 - (cat_pwpt * 1000.0))
        categories.append({"category": cat, "apt": cat_apt, "pwpt": cat_pwpt, "mqm": cat_mqm})

    sev_counts = defaultdict(int)
    for i in scorable:
        sev_counts[i.get("severity", "")] += 1

    return {
        "ewt": ewt,
        "apt": apt,
        "pwpt": pwpt,
        "mqm": mqm,
        "critical_count": crit_cnt,
        "status": status,
        "categories": categories,
        "severity_counts": dict(sev_counts),
        "issues": reported,
        "scorable_issues": scorable,
    }


def summarize_overall(findings: List[Dict[str, Any]], ewt: int) -> Dict[str, Any]:
    """Fallback programmatic overall summary/action plan across many sections."""
    by_cat_points = defaultdict(int)
    sev_counts = defaultdict(int)
    for f in findings:
        by_cat_points[f.get("type_of_issue", "")] += f.get("score_deduction", 0)
        sev_counts[f.get("severity", "")] += 1

    top_cats = sorted(by_cat_points.items(), key=lambda kv: kv[1], reverse=True)[:3]
    top_str = "; ".join(f"{c}: {p} pts" for c, p in top_cats if c)

    total_issues = len(findings)
    summ = [
        f"Aggregated across {total_issues} findings; EWT={ewt}.",
        f"Top categories by penalty points → {top_str or 'N/A'}.",
        f"Severity mix → Minor: {sev_counts.get('Minor',0)}, Major: {sev_counts.get('Major',0)}, Critical: {sev_counts.get('Critical',0)}.",
    ]
    plan = [
        "Address the top-penalty categories first with targeted checks (linters or checklists).",
        "Add a final proofing pass for brand/proper nouns and locale-specific typography before publication.",
    ]
    return {"summary": summ, "action_plan": plan}


# ────────────────────────── EXCEL REPORT WRITER ────────────────────────────
def _coerce_to_multiline(value) -> str:
    if isinstance(value, list):
        return "• " + "\n• ".join(str(x) for x in value)
    return "" if value is None else str(value)

def _fmt(ws):
    ws.page_setup.fitToWidth = 1
    for col in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
        if col in ws.column_dimensions and ws.column_dimensions[col].width is None:
            ws.column_dimensions[col].width = 18

def _write_scorecard_sheet(
    wb: Workbook,
    lang_code: str,
    metrics: Dict[str, Any],
    overall_feedback: Dict[str, Any],
    pass_threshold: int = 90,
    evaluated_scope: str = "DOCX sections",
):
    ws = wb.active
    ws.title = "Scorecard"

    bold = Font(bold=True)
    big = Font(bold=True, size=18)
    wrap = Alignment(wrap_text=True, vertical="top")
    green = PatternFill("solid", fgColor="C6EFCE")
    red = PatternFill("solid", fgColor="FFC7CE")
    light = PatternFill("solid", fgColor="F2F2F2")

    ws["A1"], ws["B1"] = "Language:", lang_code
    ws["A2"], ws["B2"] = "Evaluated Scope:", evaluated_scope
    ws["A3"], ws["B3"] = "Pass Threshold:", f"{pass_threshold}%"
    ws["A1"].font = ws["A2"].font = ws["A3"].font = bold

    ws["A5"], ws["B5"] = "Final MQM Score", round(metrics["mqm"], 2)
    ws["A5"].font = bold
    ws["B5"].font = big
    ws["A6"], ws["B6"] = "Quality Status", metrics["status"]
    ws["A6"].font = bold
    ws["B6"].fill = green if metrics["status"] == "Pass" else red
    ws["B6"].font = Font(bold=True)

    ws["A8"], ws["B8"] = "WC", metrics["ewt"]
    ws["A9"], ws["B9"] = "APT (Total Points)", metrics["apt"]
    ws["A10"], ws["B10"] = "PWPT", round(metrics["pwpt"], 3)
    # Removed duplicate "Final MQM" formula row
    ws["A12"], ws["B12"] = "Critical Issues", metrics["critical_count"]
    for r in range(8, 13):
        ws[f"A{r}"].font = bold

    ws["D8"], ws["E8"], ws["F8"], ws["G8"] = "Minor", "Major", "Critical", "Total"
    ws["D8"].font = ws["E8"].font = ws["F8"].font = ws["G8"].font = bold
    ws["D9"] = metrics["severity_counts"].get("Minor", 0)
    ws["E9"] = metrics["severity_counts"].get("Major", 0)
    ws["F9"] = metrics["severity_counts"].get("Critical", 0)
    ws["G9"] = "=SUM(D9:F9)"
    ws["D9"].fill = light; ws["E9"].fill = light; ws["F9"].fill = light; ws["G9"].fill = light

    start = 14
    ws[f"A{start}"], ws[f"B{start}"], ws[f"C{start}"], ws[f"D{start}"] = ("Category", "APT", "PWPT", "MQM")
    for c in "ABCD":
        ws[f"{c}{start}"].font = bold

    r = start + 1
    for row in metrics["categories"]:
        ws[f"A{r}"] = row["category"]
        ws[f"B{r}"] = row["apt"]
        ws[f"C{r}"] = f"=B{r}/$B$8"
        ws[f"D{r}"] = f"=MAX(0,100-(C{r}*1000))"
        r += 1

    ws["A22"], ws["A26"] = "Summary (top patterns)", "Action Plan (preventive)"
    ws["A22"].font = ws["A26"].font = bold
    ws["B22"] = _coerce_to_multiline(overall_feedback.get("summary", ""))
    ws["B26"] = _coerce_to_multiline(overall_feedback.get("action_plan", ""))
    ws["B22"].alignment = ws["B26"].alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 68
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 10
    ws.column_dimensions["G"].width = 10

    _fmt(ws)

def _write_findings_sheet(wb: Workbook, findings: List[Dict[str, Any]]):
    ws = wb.create_sheet(title="Detailed Findings")
    headers = [
        "Source File",
        "Candidate",          # renamed from "url"
        "text",
        "issue",
        "type_of_issue",
        "severity",
        "score_deduction",
        "solution",
        "updated_sentence",
    ]
    ws.append(headers)

    wrap = Alignment(wrap_text=True, vertical="top")
    for it in findings:
        ws.append(
            [
                Path(str(it.get("source_file") or "")).name,
                it.get("candidate") or it.get("url", "") or "",
                it.get("text", ""),
                it.get("issue", ""),
                it.get("type_of_issue", ""),
                it.get("severity", ""),
                it.get("score_deduction", 0),
                it.get("solution", ""),
                it.get("updated_sentence", ""),
            ]
        )

    last_row = ws.max_row
    tab = Table(displayName="Findings", ref=f"A1:I{last_row}")
    style = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False
    )
    tab.tableStyleInfo = style
    ws.add_table(tab)

    for col_letter, width in [
        ("A", 32), ("B", 48), ("C", 60), ("D", 40), ("E", 24),
        ("F", 12), ("G", 14), ("H", 40), ("I", 60)
    ]:
        ws.column_dimensions[col_letter].width = width
        for cell in ws[col_letter]:
            cell.alignment = wrap

    _fmt(ws)

def write_language_report(
    lang_code: str,
    evaluated_scope_label: str,
    evaluated_text: str,
    issues_with_url: List[Dict[str, Any]],
    output_dir: Path,
    prefix: str = "LQA_Report",
    pass_threshold: int = 90,
    save_raw_json: bool = True,
    overall_feedback_override: Optional[Dict[str, Any]] = None,
    extra_dump: Optional[Dict[str, Any]] = None,
):
    """
    Build one XLSX per language with two tabs: Scorecard & Detailed Findings.
    If overall_feedback_override is provided (from Agent 3), it's used verbatim.
    """
    ewt = compute_ewt(evaluated_text)
    metrics = aggregate_metrics(issues_with_url, ewt, pass_threshold=pass_threshold)

    if overall_feedback_override and isinstance(overall_feedback_override, dict):
        overall_feedback = overall_feedback_override
    else:
        overall_feedback = summarize_overall(metrics["scorable_issues"], ewt)

    ts = datetime.now(ZoneInfo("Africa/Cairo")).strftime("%Y%m%d_%H%M%S")
    report_path = output_dir / f"{prefix}_{lang_code}_{ts}.xlsx"

    wb = Workbook()
    _write_scorecard_sheet(
        wb,
        lang_code,
        metrics,
        overall_feedback,
        pass_threshold=pass_threshold,
        evaluated_scope=evaluated_scope_label,
    )
    _write_findings_sheet(wb, metrics["issues"])
    wb.save(report_path)
    logging.info(f"[{lang_code}] Wrote report → {report_path}")

    if save_raw_json:
        raw_payload = {
            "review_findings": metrics["issues"],
            "overall_feedback": overall_feedback,
        }
        if isinstance(extra_dump, dict):
            raw_payload["extras"] = extra_dump
        raw_path = output_dir / f"{prefix}_{lang_code}_{ts}_final.json"
        try:
            with open(raw_path, "w", encoding="utf-8") as f:
                json.dump(raw_payload, f, ensure_ascii=False, indent=2)
            logging.info(f"[{lang_code}] Saved raw JSON → {raw_path}")
        except Exception as e:
            logging.warning(f"[{lang_code}] Could not save raw JSON: {e}")


# ─────────────────────────── UTIL: Agent 3 Rollup ──────────────────────────
def build_agent3_rollup(
    lang_code: str,
    issues: List[Dict[str, Any]],
    evaluated_text_all: str,
) -> Dict[str, Any]:
    """Compact quantitative rollup for Agent 3 input."""
    ewt = compute_ewt(evaluated_text_all)
    _, norm = _issues_for_reporting_and_scoring(issues)

    severity_counts = Counter(i.get("severity", "") for i in norm)
    by_cat_points = defaultdict(int)
    for i in norm:
        cat = i.get("type_of_issue", "")
        if cat in ALLOWED_CATEGORIES:
            by_cat_points[cat] += i.get("score_deduction", 0)

    url_counts = Counter(i.get("url", "") for i in norm if i.get("url"))
    top_urls = [{"url": u, "count": c} for u, c in url_counts.most_common(10)]

    apt = sum(i.get("score_deduction", 0) for i in norm)
    pwpt = (apt / ewt) if ewt > 0 else 0.0
    mqm = max(0.0, 100.0 - (pwpt * 1000.0))

    return {
        "lang": lang_code,
        "ewt": ewt,
        "total_findings": len(norm),
        "severity_counts": dict(severity_counts),
        "points_by_category": dict(by_cat_points),
        "apt": apt,
        "pwpt": pwpt,
        "mqm": mqm,
        "top_urls_by_issue_count": top_urls,
        "categories_considered": ALLOWED_CATEGORIES,
    }



def _to_int_or_none(value: Any) -> Optional[int]:
    if value is None:
        return None
    try:
        return int(value)
    except Exception:
        return None


# ───────────────────────── DOCX → sections / edited writer ─────────────────
def _docx_style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip()
    except Exception:
        return ""


def _docx_heading_level(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    m = re.match(r"(?i)^heading\s*(\d+)$", style_name.strip())
    if m:
        return int(m.group(1))
    # Some locales / templates
    m2 = re.match(r"(?i)^t[ií]tulo\s*(\d+)$", style_name.strip())
    if m2:
        return int(m2.group(1))
    return None


def parse_docx_sections_to_df(
    path: Path,
    drop_empty: bool = True,
    prefer_level: int = 2,
    max_sections: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """
    Parse MosAIQ-style Skyscanner DOCX:
      - Heading 1 = destination URL (url_path / candidate)
      - Heading {prefer_level} (default 2) = FAQ/SEO section heading
      - Body paragraphs under each H2 = review batch content
    """
    doc = Document(str(path))
    rows: List[Dict[str, Any]] = []
    current_url = ""
    current_h2 = ""
    body_parts: List[str] = []
    body_para_indices: List[int] = []
    section_counter = 0
    h1_counter = 0

    def _flush():
        nonlocal section_counter, body_parts, body_para_indices, current_h2
        content = "\n\n".join(p for p in body_parts if p.strip())
        if drop_empty and not content.strip():
            body_parts = []
            body_para_indices = []
            return
        if not current_h2 and not content.strip():
            body_parts = []
            body_para_indices = []
            return
        section_counter += 1
        topic = current_h2 or Path(path).stem
        candidate = current_url or topic
        rows.append(
            {
                "section_id": f"{path.stem}__{section_counter:04d}",
                "url_path": current_url,
                "heading": current_h2,
                "content": content,
                "topic_name": topic,
                "candidate_name": candidate,
                "source_file": str(path.resolve()),
                "body_para_indices": list(body_para_indices),
                "h1_index": h1_counter,
            }
        )
        body_parts = []
        body_para_indices = []

    for idx, para in enumerate(doc.paragraphs):
        style_name = _docx_style_name(para)
        level = _docx_heading_level(style_name)
        text = (para.text or "").strip()

        if level == 1:
            _flush()
            h1_counter += 1
            current_url = text
            current_h2 = ""
            continue

        if level == prefer_level:
            _flush()
            current_h2 = text
            continue

        if level is not None and level != prefer_level:
            # Other heading levels: treat as structural, flush open body
            _flush()
            current_h2 = text if level < prefer_level else current_h2
            continue

        # Body / title / normal
        if current_h2 or current_url:
            # Only accumulate body once we are inside an H1 block; prefer after H2
            if current_h2:
                body_parts.append(para.text or "")
                body_para_indices.append(idx)
            elif text and style_name.lower() not in {"title", "subtitle"}:
                # Rare: body under H1 before any H2 — keep as intro batch under synthetic heading
                pass

    _flush()

    if isinstance(max_sections, int) and max_sections >= 0:
        rows = rows[:max_sections]

    return rows


def _apply_docx_language_formatting(doc: Document, lang_code: Optional[str]) -> None:
    if lang_code != "TH":
        return
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Tahoma"
            properties = run._element.get_or_add_rPr()
            fonts = properties.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                properties.insert(0, fonts)
            for font_kind in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{font_kind}"), "Tahoma")
            language = properties.find(qn("w:lang"))
            if language is None:
                language = OxmlElement("w:lang")
                properties.append(language)
            for language_kind in ("val", "eastAsia", "bidi"):
                language.set(qn(f"w:{language_kind}"), "th-TH")


def write_edited_docx(
    original_path: Path,
    section_rows: List[Dict[str, Any]],
    corrections_by_section_id: Dict[str, str],
    output_path: Path,
    lang_code: Optional[str] = None,
) -> Path:
    """
    Rebuild an edited DOCX from the original structure.
    For each H2 section, use corrected_content when present; otherwise keep original body.
    corrected_content may contain multiple paragraphs separated by blank lines / newlines.
    """
    src = Document(str(original_path))
    out = Document()

    # Map section_id -> corrected body (normalized)
    # Also build lookup by (url_path, heading) for robustness
    by_url_heading: Dict[Tuple[str, str], str] = {}
    original_by_key: Dict[Tuple[str, str], str] = {}
    for row in section_rows:
        if Path(row.get("source_file", "")).resolve() != original_path.resolve():
            # Compare by name if resolve differs across runs
            if Path(row.get("source_file", "")).name != original_path.name:
                continue
        key = (row.get("url_path") or "", row.get("heading") or "")
        original_by_key[key] = row.get("content") or ""
        sid = row.get("section_id") or ""
        if sid in corrections_by_section_id and corrections_by_section_id[sid] is not None:
            by_url_heading[key] = corrections_by_section_id[sid]

    current_url = ""
    current_h2 = ""
    pending_body: List[str] = []

    def _emit_heading(text: str, level: int) -> None:
        style = f"Heading {level}"
        p = out.add_paragraph(text, style=style)
        return p

    def _emit_body_paragraphs(text_block: str) -> None:
        if text_block is None:
            return
        stripped = text_block.strip()
        if not stripped:
            return
        # Prefer blank-line paragraph splits; otherwise keep as one paragraph
        if re.search(r"\n\s*\n", stripped):
            chunks = [c.strip() for c in re.split(r"\n\s*\n", stripped) if c.strip()]
        else:
            chunks = [stripped]
        for chunk in chunks:
            out.add_paragraph(chunk)

    def _flush_body() -> None:
        nonlocal pending_body, current_url, current_h2
        if not current_h2:
            pending_body = []
            return
        key = (current_url, current_h2)
        if key in by_url_heading:
            body = by_url_heading[key]
        else:
            body = "\n\n".join(pending_body) if pending_body else original_by_key.get(key, "")
        _emit_body_paragraphs(body or "")
        pending_body = []

    # Preserve leading non-H1 content (Title / Document N of N)
    saw_h1 = False
    for para in src.paragraphs:
        style_name = _docx_style_name(para)
        level = _docx_heading_level(style_name)
        text = para.text or ""

        if level == 1:
            if saw_h1 or current_h2:
                _flush_body()
            saw_h1 = True
            current_url = text.strip()
            current_h2 = ""
            pending_body = []
            _emit_heading(text.strip(), 1)
            continue

        if level == 2:
            _flush_body()
            current_h2 = text.strip()
            pending_body = []
            _emit_heading(text.strip(), 2)
            continue

        if not saw_h1:
            # Preamble (Title etc.)
            style = style_name if style_name else "Normal"
            try:
                out.add_paragraph(text, style=style)
            except Exception:
                out.add_paragraph(text)
            continue

        if current_h2:
            pending_body.append(text)
        elif text.strip():
            # Under H1 before H2 — keep as normal
            out.add_paragraph(text)

    _flush_body()

    _apply_docx_language_formatting(out, lang_code)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    out.save(str(output_path))
    logging.info("Wrote edited DOCX → %s", output_path)
    return output_path


def parse_csv_sections_to_list(
    path: Path,
    drop_empty: bool = True,
    heading_tag: str = "h3",
    max_sections: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """
    Parse Skyscanner CSV exports with columns:
      - page_url
      - content_html  (HTML body with repeated <h3>…</h3><p>…</p> blocks)

    Each <h3> (+ following body until the next <h3>) becomes one review batch.
    """
    rows_out: List[Dict[str, Any]] = []
    section_counter = 0

    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        if not reader.fieldnames:
            logging.error("CSV has no header row: %s", path)
            return []
        field_map = { (f or "").strip().lower(): f for f in reader.fieldnames }
        url_key = field_map.get("page_url") or field_map.get("url") or field_map.get("url_path")
        html_key = field_map.get("content_html") or field_map.get("html") or field_map.get("content")
        if not url_key or not html_key:
            logging.error(
                "CSV %s missing required columns (need page_url + content_html). Found: %s",
                path.name,
                reader.fieldnames,
            )
            return []

        for row_idx, row in enumerate(reader, start=1):
            page_url = (row.get(url_key) or "").strip()
            content_html = row.get(html_key) or ""
            if not content_html.strip():
                continue

            soup = BeautifulSoup(content_html, _PARSER)
            headings = soup.find_all(heading_tag)
            if not headings:
                # Whole cell as one batch
                plain = soup.get_text("\n", strip=True)
                if drop_empty and not plain.strip():
                    continue
                section_counter += 1
                rows_out.append(
                    {
                        "section_id": f"{path.stem}__{section_counter:04d}",
                        "url_path": page_url,
                        "heading": page_url or path.stem,
                        "content": plain,
                        "topic_name": page_url or path.stem,
                        "candidate_name": page_url,
                        "source_file": str(path.resolve()),
                        "source_format": "csv",
                        "csv_row_index": row_idx,
                    }
                )
                if isinstance(max_sections, int) and len(rows_out) >= max_sections:
                    return rows_out
                continue

            for h in headings:
                heading_text = h.get_text(" ", strip=True)
                # Collect siblings until next same-level heading
                body_parts: List[str] = []
                for sib in h.next_siblings:
                    sib_name = getattr(sib, "name", None)
                    if sib_name == heading_tag:
                        break
                    if sib_name in {"p", "div"}:
                        t = sib.get_text(" ", strip=True)
                        if t:
                            body_parts.append(t)
                    elif sib_name in {"ul", "ol"}:
                        for li in sib.find_all("li"):
                            t = li.get_text(" ", strip=True)
                            if t:
                                body_parts.append(t)
                    elif sib_name is None:
                        t = str(sib).strip()
                        if t:
                            body_parts.append(t)
                content = "\n\n".join(body_parts)
                if drop_empty and not content.strip() and not heading_text:
                    continue

                section_counter += 1
                topic = heading_text or page_url or path.stem
                rows_out.append(
                    {
                        "section_id": f"{path.stem}__{section_counter:04d}",
                        "url_path": page_url,
                        "heading": heading_text,
                        "content": content,
                        "topic_name": topic,
                        "candidate_name": page_url,
                        "source_file": str(path.resolve()),
                        "source_format": "csv",
                        "csv_row_index": row_idx,
                    }
                )
                if isinstance(max_sections, int) and len(rows_out) >= max_sections:
                    return rows_out

    return rows_out


def _plain_body_to_html_paragraphs(text: str) -> str:
    """Convert corrected plain text (blank-line separated) into one or more <p> tags."""
    stripped = (text or "").strip()
    if not stripped:
        return "<p></p>"
    if re.search(r"\n\s*\n", stripped):
        chunks = [c.strip() for c in re.split(r"\n\s*\n", stripped) if c.strip()]
    else:
        chunks = [stripped]
    return "".join(f"<p>{html_escape(c)}</p>" for c in chunks)


def write_edited_csv(
    original_path: Path,
    section_rows: List[Dict[str, Any]],
    corrections_by_section_id: Dict[str, str],
    output_path: Path,
    heading_tag: str = "h3",
) -> Path:
    """
    Rebuild an edited CSV with the same page_url / content_html schema.
    Each page's H3 sections are reconstructed from Agent 2 corrected_content.
    """
    # Group section rows belonging to this file, preserving parse order
    page_sections: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for row in section_rows:
        src = Path(row.get("source_file", ""))
        if src.resolve() != original_path.resolve() and src.name != original_path.name:
            continue
        page_sections[row.get("url_path") or ""].append(row)

    # Read original order of URLs
    original_urls: List[str] = []
    with original_path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        field_map = { (f or "").strip().lower(): f for f in (reader.fieldnames or []) }
        url_key = field_map.get("page_url") or field_map.get("url") or field_map.get("url_path")
        for row in reader:
            original_urls.append((row.get(url_key) or "").strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=["page_url", "content_html"], quoting=csv.QUOTE_MINIMAL)
        writer.writeheader()
        seen = set()
        for page_url in original_urls:
            if page_url in seen:
                # Duplicate URL rows: still write using same sections
                pass
            seen.add(page_url)
            sections = page_sections.get(page_url) or []
            html_parts: List[str] = []
            for sec in sections:
                sid = sec.get("section_id") or ""
                heading = sec.get("heading") or ""
                body = corrections_by_section_id.get(sid)
                if body is None:
                    body = sec.get("content") or ""
                html_parts.append(
                    f"<{heading_tag}>{html_escape(heading)}</{heading_tag}>"
                    f"{_plain_body_to_html_paragraphs(body)}"
                )
            writer.writerow({"page_url": page_url, "content_html": "".join(html_parts)})

    logging.info("Wrote edited CSV → %s", output_path)
    return output_path


def _build_corrected_content_fallback(original_content: str, findings: List[Dict[str, Any]]) -> str:
    """
    If Agent 2 omits corrected_content, apply updated_sentence replacements
    to the original section body (best-effort).
    """
    text = original_content or ""
    for f in findings or []:
        src = (f.get("text") or "").strip()
        upd = (f.get("updated_sentence") or "").strip()
        if src and upd and src in text:
            text = text.replace(src, upd, 1)
    return text


_PLACE_NAME_RE = re.compile(r"(?i)place name|city name|geographic name|location name")
_LOCALIZATION_RE = re.compile(
    r"(?i)transliterat|locali[sz]|latin (?:alphabet|script)|written in greek|greek spelling"
)
_LATIN_LETTER_RE = re.compile(r"[A-Za-z\u00c0-\u00d6\u00d8-\u00f6\u00f8-\u00ff]")


def apply_project_finding_rules(
    lang_code: str,
    original_content: str,
    findings: List[Dict[str, Any]],
    corrected_content: Any,
) -> Tuple[List[Dict[str, Any]], str, List[Dict[str, Any]]]:
    if lang_code != "GR":
        corrected = corrected_content if isinstance(corrected_content, str) and corrected_content.strip() else (
            _build_corrected_content_fallback(original_content, findings) if findings else original_content
        )
        return findings, corrected, []

    removed = []
    kept = []
    for finding in findings:
        explanation = f"{finding.get('issue', '')} {finding.get('solution', '')}"
        evidence = f"{finding.get('text', '')} {finding.get('updated_sentence', '')}"
        if (
            _PLACE_NAME_RE.search(explanation)
            and _LOCALIZATION_RE.search(explanation)
            and _LATIN_LETTER_RE.search(evidence)
        ):
            removed.append(finding)
        else:
            kept.append(finding)

    if removed:
        corrected = _build_corrected_content_fallback(original_content, kept)
    elif isinstance(corrected_content, str) and corrected_content.strip():
        corrected = corrected_content
    else:
        corrected = _build_corrected_content_fallback(original_content, kept) if kept else original_content
    return kept, corrected, removed


# ─────────────────────────── MAIN PIPELINE ─────────────────────────────────
def _build_iso_to_short_map(lang_map: Dict[str, str]) -> Dict[str, str]:
    iso_to_short: Dict[str, str] = {}
    for short_code, locale_code in (lang_map or {}).items():
        iso = str(locale_code).split("-", 1)[0].strip().lower()
        if iso and iso not in iso_to_short:
            iso_to_short[iso] = short_code
    return iso_to_short


def _infer_lang_code_for_input_file(input_file: Path, lang_map: Dict[str, str]) -> Optional[str]:
    """
    Support multiple naming conventions:
    - Legacy: `FR_1.html` / `FR_1.docx` where prefix is a configured short code.
    - MosAIQ / Destination-Hub: `...-br-flights-to-region-01.docx` where `_br_` / `-br-` is short or ISO.
    - Browser-saved HTML: embedded `<!-- saved from url=..._fr_... -->`.
    """
    name_upper = input_file.name.upper()
    for short_code in (lang_map or {}).keys():
        if name_upper.startswith(f"{short_code.upper()}_") or name_upper.startswith(f"{short_code.upper()}-"):
            return short_code

    # Short codes as delimited tokens in the stem (e.g. -br-, _DE_)
    short_keys = sorted((lang_map or {}).keys(), key=len, reverse=True)
    if short_keys:
        short_pat = re.compile(
            rf"(?i)(?:^|[_-])(?P<code>{'|'.join(map(re.escape, short_keys))})(?:$|[_-])"
        )
        m_short = short_pat.search(input_file.stem)
        if m_short:
            code = m_short.group("code").upper()
            for short_code in lang_map.keys():
                if short_code.upper() == code:
                    return short_code

    iso_to_short = _build_iso_to_short_map(lang_map)
    if not iso_to_short:
        return None

    iso_keys = sorted(iso_to_short.keys(), key=len, reverse=True)
    iso_pat = re.compile(rf"(?i)(?:^|[_-])(?P<iso>{'|'.join(map(re.escape, iso_keys))})(?:$|[_-])")

    m = iso_pat.search(input_file.stem)
    if m:
        return iso_to_short.get(m.group("iso").lower())

    # HTML-only fallback: embedded "saved from url" comment
    if input_file.suffix.lower() in {".html", ".htm"}:
        try:
            html = input_file.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return None
        candidate = _fallback_candidate_path_from_html(html, input_file)
        m2 = iso_pat.search(candidate)
        if m2:
            return iso_to_short.get(m2.group("iso").lower())

    return None


def _normalize_requested_language(
    requested_language: Optional[str], lang_map: Dict[str, str]
) -> Optional[str]:
    if not requested_language:
        return None
    requested = requested_language.strip().upper()
    for code in lang_map:
        if code.upper() == requested:
            return code
    raise ValueError(f"Unknown language code: {requested}")


def discover_input_files(
    input_dir: Path,
    input_format: str,
    lang_map: Dict[str, str],
    requested_language: Optional[str] = None,
) -> Dict[str, List[Path]]:
    files_by_lang: Dict[str, List[Path]] = defaultdict(list)
    extensions = {
        "csv": ("*.csv",),
        "docx": ("*.docx",),
        "html": ("*.html",),
        "auto": ("*.csv", "*.docx", "*.html"),
    }
    for pattern in extensions.get(input_format, ()):
        for input_file in sorted(input_dir.glob(pattern)):
            lang_code = _infer_lang_code_for_input_file(input_file, lang_map)
            if not lang_code:
                logging.warning("Skipping %s: could not infer language code from filename.", input_file.name)
            elif requested_language is None or lang_code == requested_language:
                files_by_lang[lang_code].append(input_file)
    return files_by_lang


def render_system_prompt(
    template: str,
    lang_code: str,
    locale_instructions: str,
    project_instructions: str,
) -> str:
    return (
        template.replace("$lang", lang_code)
        .replace("$locale_instructions", locale_instructions)
        .replace("$project_instructions", project_instructions)
    )


def write_agent_audit(
    output_path: Path,
    lang_code: str,
    section_rows: List[Dict[str, Any]],
    agent1_results: List[Dict[str, Any]],
    agent2_results: List[Dict[str, Any]],
) -> Path:
    def count_findings(results: List[Dict[str, Any]]) -> int:
        return sum(
            len(result.get("review_findings", []))
            for result in results
            if isinstance(result, dict) and isinstance(result.get("review_findings"), list)
        )

    payload = {
        "language": lang_code,
        "agent1_finding_count": count_findings(agent1_results),
        "agent2_finding_count": count_findings(agent2_results),
        "sections": [
            {
                "section_id": row.get("section_id"),
                "source_file": Path(str(row.get("source_file", ""))).name,
                "agent1": agent1_results[index],
                "agent2": agent2_results[index],
            }
            for index, row in enumerate(section_rows)
        ],
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    logging.info("[%s] Saved Agent 1/2 audit -> %s", lang_code, output_path)
    return output_path


async def main(selected_language: Optional[str] = None):
    """Three-agent, windowed per-language processing for DOCX sections."""
    file_handler: Optional[logging.FileHandler] = None
    log_file_path: Optional[Path] = None
    try:
        config = load_config()
        requested_language = _normalize_requested_language(
            selected_language, config.get("lang_map", {}) or {}
        )
        api_key = resolve_api_key(config)

        # Paths
        script_dir = Path(__file__).parent
        input_dir = script_dir / config["paths"]["input"]
        prompts_dir = script_dir / config["paths"]["prompts"]
        ref_dir = script_dir / config["paths"]["reference"]
        output_dir = script_dir / config["paths"]["output"]
        output_dir.mkdir(exist_ok=True)

        # Add file logging now that we know output_dir
        ts_run = datetime.now(ZoneInfo("Africa/Cairo")).strftime("%Y%m%d_%H%M%S")
        log_scope = f"_{requested_language}" if requested_language else ""
        log_file_path = output_dir / f"lqa_run{log_scope}_{ts_run}.log"
        file_handler = logging.FileHandler(log_file_path, encoding="utf-8")
        file_handler.setLevel(logging.INFO)
        file_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
        logging.getLogger().addHandler(file_handler)
        logger.info("File logging enabled → %s", log_file_path)
        if requested_language:
            logger.info("Language filter active → %s", requested_language)

        # Load prompts (Agent 1 & 2)
        with open(prompts_dir / config["file_names"]["system_prompt"], "r", encoding="utf-8") as f:
            system_prompt_template_1 = f.read()
        with open(prompts_dir / config["file_names"]["user_prompt"], "r", encoding="utf-8") as f:
            user_prompt_template_1 = f.read()

        with open(prompts_dir / config["file_names"]["system_prompt_2"], "r", encoding="utf-8") as f:
            system_prompt_template_2 = f.read()
        with open(prompts_dir / config["file_names"]["user_prompt_2"], "r", encoding="utf-8") as f:
            user_prompt_template_2 = f.read()
        project_prompt_name = config["file_names"].get("project_instructions")
        project_instructions = (
            (prompts_dir / project_prompt_name).read_text(encoding="utf-8")
            if project_prompt_name
            else ""
        )

        # Load prompts (Agent 3)
        with open(prompts_dir / config["file_names"]["system_prompt_3"], "r", encoding="utf-8") as f:
            system_prompt_3 = f.read()
        with open(prompts_dir / config["file_names"]["user_prompt_3"], "r", encoding="utf-8") as f:
            user_prompt_template_3 = f.read()

        # Load glossary/context
        glossary_terms = load_glossary(ref_dir, config["file_names"]["glossary"], config["lang_map"])

        # Batching config
        batching = config.get("batching", {})
        limits = config.get("limits", {}) or {}
        max_batches_per_window = int(batching.get("max_batches_per_window", 50))
        window_pause_seconds = float(batching.get("window_pause_seconds", 10))
        combine_reports = bool(batching.get("combine_reports_per_language", True))
        drop_empty_sections = bool(batching.get("drop_empty_sections", True))
        max_hr_sections_per_language = _to_int_or_none(limits.get("max_hr_sections_per_language"))
        max_final_batches_per_language = _to_int_or_none(limits.get("max_final_batches_per_language"))
        legacy_max_sections = _to_int_or_none(batching.get("max_sections_per_language"))
        if max_final_batches_per_language is None:
            max_final_batches_per_language = legacy_max_sections
        max_words_per_language = limits.get("max_words_per_language")
        if max_words_per_language is None:
            max_words_per_language = batching.get("max_words_per_language", None)
        max_words_per_language = _to_int_or_none(max_words_per_language)

        # Parsing mode config
        parsing_cfg = config.get("parsing", {}) or {}
        split_mode = str(parsing_cfg.get("split_mode", "heading")).lower()
        heading_level_preference = int(parsing_cfg.get("heading_level_preference", 2))
        input_format = str(parsing_cfg.get("input_format", "csv")).lower()  # csv | docx | html | auto
        csv_heading_tag = str(parsing_cfg.get("csv_heading_tag", "h3")).lower()

        prefix = config["file_names"]["output_report_prefix"]
        edited_prefix = config["file_names"].get("edited_output_prefix", "Edited")

        # Discover inputs grouped by language
        lang_map_cfg: Dict[str, str] = config.get("lang_map", {}) or {}
        files_by_lang = discover_input_files(
            input_dir, input_format, lang_map_cfg, requested_language
        )

        if not files_by_lang:
            logging.info("No CSV/DOCX/HTML input files found. Exiting.")
            return

        # Process languages sequentially
        for lang_code in sorted(files_by_lang.keys()):
            logging.info(f"===== Language {lang_code}: parsing input files → sections/batches =====")

            try:
                with open(prompts_dir / f"{lang_code}_instructions.txt", "r", encoding="utf-8") as f:
                    locale_instructions = f.read()
            except FileNotFoundError:
                logging.warning(f"No instructions file for {lang_code}. Using empty instructions.")
                locale_instructions = "No language-specific instructions provided."

            system_prompt_1 = render_system_prompt(
                system_prompt_template_1,
                lang_code,
                locale_instructions,
                project_instructions,
            )
            system_prompt_2 = render_system_prompt(
                system_prompt_template_2,
                lang_code,
                locale_instructions,
                project_instructions,
            )

            terms = glossary_terms.get(lang_code, "N/A")

            section_rows: List[Dict[str, Any]] = []
            remaining_sections = max_hr_sections_per_language
            for path in files_by_lang[lang_code]:
                per_file_cap: Optional[int] = None
                if isinstance(remaining_sections, int):
                    if remaining_sections <= 0:
                        break
                    per_file_cap = remaining_sections

                if path.suffix.lower() == ".csv":
                    parsed_rows = parse_csv_sections_to_list(
                        path,
                        drop_empty=drop_empty_sections,
                        heading_tag=csv_heading_tag,
                        max_sections=per_file_cap,
                    )
                elif path.suffix.lower() == ".docx":
                    parsed_rows = parse_docx_sections_to_df(
                        path,
                        drop_empty=drop_empty_sections,
                        prefer_level=heading_level_preference,
                        max_sections=per_file_cap,
                    )
                else:
                    # Legacy HTML path retained for compatibility
                    include_parent_h1 = str(parsing_cfg.get("include_parent_h1", "prefix_all")).lower()
                    create_intro_batch_if_long = bool(parsing_cfg.get("create_intro_batch_if_long", True))
                    min_intro_chars = int(parsing_cfg.get("min_intro_chars", 200))
                    allow_fallback_to_h1 = bool(parsing_cfg.get("allow_fallback_to_h1", False))
                    if split_mode == "heading":
                        parsed_rows = parse_html_sections_to_df_heading(
                            path,
                            drop_empty=drop_empty_sections,
                            prefer_level=heading_level_preference,
                            include_parent_h1=include_parent_h1,
                            create_intro_batch_if_long=create_intro_batch_if_long,
                            min_intro_chars=min_intro_chars,
                            allow_fallback_to_h1=allow_fallback_to_h1,
                            max_hr_sections=per_file_cap,
                        )
                    else:
                        parsed_rows = parse_html_sections_to_df(
                            path,
                            drop_empty=drop_empty_sections,
                            max_hr_sections=per_file_cap,
                        )

                if not combine_reports:
                    pass
                section_rows.extend(parsed_rows)

                if isinstance(remaining_sections, int):
                    remaining_sections -= len(parsed_rows)

            if not section_rows:
                logging.warning(f"[{lang_code}] No content sections/batches found after parsing. Skipping.")
                continue

            if isinstance(max_final_batches_per_language, int):
                section_rows = section_rows[:max_final_batches_per_language]
                logging.info(
                    f"[{lang_code}] Applied max_final_batches_per_language={max_final_batches_per_language}; "
                    f"remaining batches: {len(section_rows)}"
                )

            if isinstance(max_words_per_language, int):
                capped_rows: List[Dict[str, Any]] = []
                total_words = 0
                limit = max_words_per_language
                limit_with_buffer = int(limit * 1.05)
                for row in section_rows:
                    words = compute_ewt(row.get("content", "") or "")
                    projected = total_words + words
                    if projected <= limit_with_buffer:
                        capped_rows.append(row)
                        total_words = projected
                        if projected >= limit_with_buffer:
                            break
                    else:
                        break
                if len(capped_rows) < len(section_rows):
                    logging.info(
                        f"[{lang_code}] Applied max_words_per_language={limit} (+5% buffer={limit_with_buffer}); "
                        f"kept {len(capped_rows)} sections totaling ~{total_words} words; "
                        f"dropped {len(section_rows)-len(capped_rows)} sections."
                    )
                section_rows = capped_rows

            total_sections = len(section_rows)
            logging.info(f"[{lang_code}] Total batches to process: {total_sections}")
            lang_output_dir = output_dir / lang_code
            lang_output_dir.mkdir(parents=True, exist_ok=True)

            # ── PASS 1: Agent 1
            agent1_results: List[Dict[str, Any]] = [None] * total_sections  # type: ignore

            all_prompts_1: List[Dict[str, str]] = []
            logged_a1_prompt = False
            for row in section_rows:
                content = row["content"]
                topic_name = row.get("topic_name") or row.get("candidate_name") or "N/A"
                user_prompt_1 = (
                    user_prompt_template_1
                    .replace("$content", content)
                    .replace("$terms", terms)
                    .replace("$topic_name", topic_name)
                )
                if not logged_a1_prompt:
                    logger.info(
                        "[A1][%s][%s] SYSTEM PROMPT (first batch only):\n%s",
                        lang_code,
                        row.get("section_id", "?"),
                        system_prompt_1,
                    )
                    logger.info(
                        "[A1][%s][%s] USER PROMPT (first batch only):\n%s",
                        lang_code,
                        row.get("section_id", "?"),
                        user_prompt_1,
                    )
                    logged_a1_prompt = True
                all_prompts_1.append({"system": system_prompt_1, "prompt": user_prompt_1})

            for start in range(0, total_sections, max_batches_per_window):
                end = min(start + max_batches_per_window, total_sections)
                window_prompts = all_prompts_1[start:end]
                logging.info(
                    f"[{lang_code}] Agent 1 window {start}-{end-1} ({len(window_prompts)} batches)"
                )
                window_results = await get_gemini_lqa_batch(
                    api_key=api_key,
                    llm_model=config["llm_settings"]["model"],
                    temp=config["llm_settings"]["temperature"],
                    prompt_parts_list=window_prompts,
                )
                agent1_results[start:end] = window_results
                if end < total_sections:
                    logging.info(
                        f"[{lang_code}] Waiting {window_pause_seconds}s before next Agent 1 window…"
                    )
                    await asyncio.sleep(window_pause_seconds)

            # ── PASS 2: Agent 2
            agent2_results: List[Dict[str, Any]] = [None] * total_sections  # type: ignore
            all_prompts_2: List[Dict[str, str]] = []
            logged_a2_prompt = False
            for idx, row in enumerate(section_rows):
                content = row["content"]
                topic_name = row.get("topic_name") or row.get("candidate_name") or "N/A"
                draft = agent1_results[idx]
                if not isinstance(draft, dict) or "error" in draft:
                    draft_obj = {
                        "review_findings": [],
                        "overall_feedback": {
                            "summary": "Draft unavailable due to error.",
                            "action_plan": "",
                        },
                    }
                else:
                    draft_obj = draft
                draft_json_str = json.dumps(draft_obj, ensure_ascii=False, indent=2)
                user_prompt_2 = (
                    user_prompt_template_2
                    .replace("$content", content)
                    .replace("$terms", terms)
                    .replace("$draft_report", draft_json_str)
                    .replace("$topic_name", topic_name)
                )
                if not logged_a2_prompt:
                    logger.info(
                        "[A2][%s][%s] SYSTEM PROMPT (first batch only):\n%s",
                        lang_code,
                        section_rows[idx].get("section_id", "?"),
                        system_prompt_2,
                    )
                    logger.info(
                        "[A2][%s][%s] USER PROMPT (first batch only):\n%s",
                        lang_code,
                        section_rows[idx].get("section_id", "?"),
                        user_prompt_2,
                    )
                    logged_a2_prompt = True
                all_prompts_2.append({"system": system_prompt_2, "prompt": user_prompt_2})

            for start in range(0, total_sections, max_batches_per_window):
                end = min(start + max_batches_per_window, total_sections)
                window_prompts = all_prompts_2[start:end]
                logging.info(
                    f"[{lang_code}] Agent 2 window {start}-{end-1} ({len(window_prompts)} batches)"
                )
                window_results = await get_gemini_lqa_batch(
                    api_key=api_key,
                    llm_model=config["llm_settings"]["model"],
                    temp=config["llm_settings"]["temperature"],
                    prompt_parts_list=window_prompts,
                )
                agent2_results[start:end] = window_results
                if end < total_sections:
                    logging.info(
                        f"[{lang_code}] Waiting {window_pause_seconds}s before next Agent 2 window…"
                    )
                    await asyncio.sleep(window_pause_seconds)

            failed_indices = [
                i
                for i, res in enumerate(agent2_results)
                if (not isinstance(res, dict)) or ("error" in res)
            ]
            if failed_indices:
                logging.warning(
                    f"[{lang_code}] Retrying {len(failed_indices)} Agent 2 batches that failed."
                )
                for idx in failed_indices:
                    try:
                        retry_result = await get_gemini_lqa_batch(
                            api_key=api_key,
                            llm_model=config["llm_settings"]["model"],
                            temp=config["llm_settings"]["temperature"],
                            prompt_parts_list=[all_prompts_2[idx]],
                        )
                        agent2_results[idx] = (
                            retry_result[0] if retry_result else {"error": "Empty retry result"}
                        )
                    except Exception as exc:  # noqa: BLE001
                        agent2_results[idx] = {"error": f"Retry failed: {exc}"}

            audit_ts = datetime.now(ZoneInfo("Africa/Cairo")).strftime("%Y%m%d_%H%M%S")
            write_agent_audit(
                lang_output_dir / f"Agent_Audit_{lang_code}_{audit_ts}.json",
                lang_code,
                section_rows,
                agent1_results,
                agent2_results,
            )

            # Consolidate findings + collect corrected_content for edited DOCX
            consolidated_issues: List[Dict[str, Any]] = []
            corrections_by_section_id: Dict[str, str] = {}
            for idx, row in enumerate(section_rows):
                url = row.get("url_path") or ""
                candidate = row.get("candidate_name") or row.get("topic_name") or ""
                source_file = Path(str(row.get("source_file") or "")).name
                sid = row.get("section_id") or f"row_{idx}"
                res = agent2_results[idx]
                if not isinstance(res, dict):
                    consolidated_issues.append(
                        {
                            "source_file": source_file,
                            "url": url,
                            "candidate": candidate,
                            "text": "",
                            "issue": "Agent 2 returned an unexpected payload for this batch.",
                            "type_of_issue": "Accuracy",
                            "solution": "",
                            "updated_sentence": "",
                            "severity": "Major",
                            "score_deduction": 0,
                        }
                    )
                    corrections_by_section_id[sid] = row.get("content") or ""
                    continue
                if "error" in res:
                    consolidated_issues.append(
                        {
                            "source_file": source_file,
                            "url": url,
                            "candidate": candidate,
                            "text": "",
                            "issue": f"Agent 2 error: {res['error']}",
                            "type_of_issue": "Accuracy",
                            "solution": "",
                            "updated_sentence": "",
                            "severity": "Major",
                            "score_deduction": 0,
                        }
                    )
                    corrections_by_section_id[sid] = row.get("content") or ""
                    continue

                findings = res.get("review_findings", []) or []
                findings, corrected, removed = apply_project_finding_rules(
                    lang_code,
                    row.get("content") or "",
                    findings,
                    res.get("corrected_content"),
                )
                if removed:
                    logging.info(
                        "[%s][%s] Removed %d excluded Latin-place-name finding(s).",
                        lang_code,
                        sid,
                        len(removed),
                    )
                for f in findings:
                    f = dict(f)
                    f["source_file"] = source_file
                    f["url"] = url
                    f["candidate"] = candidate
                    f.pop("question", None)
                    consolidated_issues.append(f)

                corrections_by_section_id[sid] = corrected

            evaluated_text_all = "\n\n".join(
                r["content"] for r in section_rows if r.get("content")
            )

            # ── Agent 3
            rollup = build_agent3_rollup(lang_code, consolidated_issues, evaluated_text_all)
            rollup_json = json.dumps(rollup, ensure_ascii=False, indent=2)
            issues_json = json.dumps(consolidated_issues, ensure_ascii=False, indent=2)

            user_prompt_3 = (
                user_prompt_template_3
                .replace("$lang", lang_code)
                .replace("$rollup", rollup_json)
                .replace("$final_issues", issues_json)
            )
            agent3_parts = [{"system": system_prompt_3, "prompt": user_prompt_3}]
            logging.info(f"[{lang_code}] Agent 3 (consolidator) - starting")
            logger.info("[A3][%s] SYSTEM PROMPT (first batch only):\n%s", lang_code, system_prompt_3)
            logger.info("[A3][%s] USER PROMPT (first batch only):\n%s", lang_code, user_prompt_3)

            agent3_result = await get_gemini_lqa_batch(
                api_key=api_key,
                llm_model=config["llm_settings"]["model"],
                temp=config["llm_settings"]["temperature"],
                prompt_parts_list=agent3_parts,
            )
            agent3_output = agent3_result[0] if agent3_result else {}

            overall_feedback_agent3: Optional[Dict[str, Any]] = None
            if isinstance(agent3_output, dict) and not agent3_output.get("error"):
                ofb = agent3_output.get("overall_feedback")
                if isinstance(ofb, dict) and "summary" in ofb and "action_plan" in ofb:
                    overall_feedback_agent3 = {
                        "summary": ofb.get("summary", []),
                        "action_plan": ofb.get("action_plan", []),
                    }
                else:
                    logging.warning(
                        f"[{lang_code}] Agent 3 output missing 'overall_feedback', using fallback."
                    )
            else:
                logging.warning(
                    f"[{lang_code}] Agent 3 error or invalid payload, using fallback: {agent3_output}"
                )

            source_suffixes = {
                Path(r["source_file"]).suffix.lower()
                for r in section_rows
                if r.get("source_file")
            }
            if ".csv" in source_suffixes:
                scope_kind = f"CSV <{csv_heading_tag}> batches"
            elif ".docx" in source_suffixes:
                scope_kind = "DOCX heading-based batches"
            else:
                scope_kind = "HTML batches"
            scope_label = (
                f"{scope_kind} ({total_sections} batches; window={max_batches_per_window})"
            )
            write_language_report(
                lang_code=lang_code,
                evaluated_scope_label=scope_label,
                evaluated_text=evaluated_text_all,
                issues_with_url=consolidated_issues,
                output_dir=lang_output_dir,
                prefix=prefix,
                pass_threshold=int(config.get("scoring", {}).get("pass_threshold", 90)),
                save_raw_json=True,
                overall_feedback_override=overall_feedback_agent3,
                extra_dump={"agent3_rollup": rollup, "agent3_raw": agent3_output},
            )

            # Write edited deliverable per source file (CSV and/or DOCX)
            ts_edit = datetime.now(ZoneInfo("Africa/Cairo")).strftime("%Y%m%d_%H%M%S")
            source_files = sorted(
                {
                    Path(r["source_file"])
                    for r in section_rows
                    if r.get("source_file")
                    and Path(r["source_file"]).suffix.lower() in {".docx", ".csv"}
                },
                key=lambda p: p.name,
            )
            for src_path in source_files:
                if not src_path.exists():
                    alt = input_dir / src_path.name
                    if alt.exists():
                        src_path = alt
                    else:
                        logging.warning("[%s] Source file missing: %s", lang_code, src_path)
                        continue
                suffix = src_path.suffix.lower()
                out_name = f"{edited_prefix}_{lang_code}_{src_path.stem}_{ts_edit}{suffix}"
                out_path = lang_output_dir / out_name
                try:
                    if suffix == ".csv":
                        write_edited_csv(
                            original_path=src_path,
                            section_rows=section_rows,
                            corrections_by_section_id=corrections_by_section_id,
                            output_path=out_path,
                            heading_tag=csv_heading_tag,
                        )
                    else:
                        write_edited_docx(
                            original_path=src_path,
                            section_rows=section_rows,
                            corrections_by_section_id=corrections_by_section_id,
                            output_path=out_path,
                            lang_code=lang_code,
                        )
                except Exception as exc:  # noqa: BLE001
                    logging.error(
                        "[%s] Failed to write edited file for %s: %s",
                        lang_code,
                        src_path.name,
                        exc,
                        exc_info=True,
                    )

            logging.info(f"[{lang_code}] Completed all three agents.\n")

        logging.info("LQA pipeline finished successfully (3-agent).")
        if log_file_path:
            logging.info("Full run log saved to %s", log_file_path)

    finally:
        if file_handler is not None:
            try:
                file_handler.flush()
            finally:
                file_handler.close()
                logging.getLogger().removeHandler(file_handler)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run the Skyscanner LQA pipeline.")
    parser.add_argument("--language", help="Process one configured short code, for example GR.")
    args = parser.parse_args()
    asyncio.run(main(args.language))
